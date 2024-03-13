from docx import Document
from io import BytesIO
import re
from fastapi import UploadFile
from wix_schemas import *
from wix_api import AsyncClient, WixBlogRoute


async def create_single_draft(file: UploadFile):
    if file.filename.endswith(".docx"):
        file_name = file.filename.replace(".docx", "")
        file_name = file_name[13:]
        file_content = await file.read()
        document = Document(BytesIO(file_content))  # Создаем объект документа из байтов
        new_content = extract_text_and_tables(document)
        new_draft = DraftPost(
            title=file_name,
            rich_content=RichContent(
                nodes=map_to_rich_content(
                    parse_content(new_content)
                )
            )
        )
        wix_api_client = AsyncClient()
        created_draft = await wix_api_client.make_request(
            new_draft, WixBlogRoute.create_draft_post
        )
        return created_draft
    else:
        return


def extract_text_and_tables(doc: Document):
    full_text = []

    for element in doc.element.body:
        if element.tag.endswith('}p'):
            # Получаем индекс и преобразуем его в int
            index = int(element.xpath('count(preceding-sibling::w:p)'))
            paragraph = doc.paragraphs[index]
            full_text.append(paragraph.text)
        elif element.tag.endswith('}tbl'):
            index = int(element.xpath('count(preceding-sibling::w:tbl)'))
            table = doc.tables[index]
            # Начинаем формировать HTML-таблицу
            table_html = "<table border='1'>"
            for row in table.rows:
                table_html += "<tr>"
                for cell in row.cells:
                    table_html += f"<td>{cell.text}</td>"
                table_html += "</tr>"
            table_html += "</table>"
            full_text.append(table_html)

    return ' '.join(full_text)


async def create_single_draft_test():
    new_draft = DraftPost(
        title="test title",
        rich_content=RichContent(
            nodes=map_to_rich_content(
                [{"type": "HEADING", "text": "heading text"},
                 {"type": "TEXT", "text": "text text"}]
            )
        )
    )
    wix_api_client = AsyncClient()
    created_draft = await wix_api_client.make_request(
        new_draft, WixBlogRoute.create_draft_post
    )
    return created_draft.json()


def parse_content(input_text):
    sections = []
    regex = re.compile(r'(HHH|PPP|TTT)(.*?)\1', re.DOTALL)

    for match in regex.finditer(input_text):
        type_tag = match.group(1)
        text = match.group(2).strip()
        if text.endswith('/'):
            text = text[:-1]
        # Определение типа на основе тега
        type_ = None
        if type_tag == 'HHH':
            type_ = 'HEADING'
        elif type_tag == 'PPP':
            type_ = 'TEXT'
        elif type_tag == 'TTT':
            type_ = 'TABLE'

        sections.append({'type': type_, 'text': text})

    return sections


def map_to_rich_content(sections: List[Dict]):
    nodes = []
    for section in sections:
        section_type = section["type"]
        text_node = Node(
            type="TEXT",
            textData=TextData(
                text=section["text"]
            )
        )
        empty_node = Node(
            type="PARAGRAPH"
        )
        if section_type == "HEADING":
            node_for_section = Node(
                type="HEADING",
                nodes=[
                    text_node
                ],
                headingData=HeadingData(
                    level=2
                )
            )
        else:
            node_for_section = Node(
                type="PARAGRAPH",
                nodes=[
                    text_node
                ],
                paragraphData={}
            )
        nodes.append(node_for_section)
        nodes.append(empty_node)

    return nodes

